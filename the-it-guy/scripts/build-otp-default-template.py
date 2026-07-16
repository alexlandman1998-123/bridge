#!/usr/bin/env python3
from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("assets/legal-templates/otp_default_v1.docx")
INK = RGBColor(16, 34, 54)
MUTED = RGBColor(91, 101, 114)
ACCENT = RGBColor(28, 90, 120)


def set_font(run, size=10.5, bold=False, color=INK):
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def field(doc, label, placeholder):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True)
    r = p.add_run(f"{{{placeholder}}}")
    set_font(r)
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, size=15 if level == 1 else 12, bold=True, color=ACCENT if level == 1 else INK)
    return p


def signature_block(doc, label, name_key):
    heading(doc, label, 2)
    field(doc, "Name", name_key)
    body(doc, "Signature: __________________________________________")
    body(doc, "Initials: ________________    Date: __________________")


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.75)
section.right_margin = Inches(0.8)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(0.8)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

normal = doc.styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("OFFER TO PURCHASE | CONTROLLED LEGAL TEMPLATE"), size=8, bold=True, color=MUTED)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("Transaction {transaction_reference}"), size=8, color=MUTED)

# Page 1 - customer-pack opening and parties.
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(30)
title.paragraph_format.space_after = Pt(4)
set_font(title.add_run("OFFER TO PURCHASE"), size=25, bold=True, color=INK)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(20)
set_font(subtitle.add_run("Residential immovable property"), size=13, color=ACCENT)
field(doc, "Property", "property_address")
field(doc, "Agency", "organisation_name")
field(doc, "Agent", "agent_full_name")
field(doc, "FFC number", "agent_ffc_number")
body(doc, "This Offer to Purchase becomes a deed of sale when accepted by the Seller in writing. The schedules, standard terms, special conditions and annexures form one agreement.")
heading(doc, "1. Parties")
field(doc, "Purchaser", "buyer_full_name")
field(doc, "Identity / registration number", "buyer_id_number")
field(doc, "Entity type", "buyer_entity_type")
field(doc, "Email", "buyer_email")
field(doc, "Telephone", "buyer_phone")
field(doc, "Seller", "seller_full_name")
field(doc, "Identity / registration number", "seller_id_number")
field(doc, "Entity type", "seller_entity_type")
field(doc, "Email", "seller_email")
field(doc, "Telephone", "seller_phone")

doc.add_page_break()
# Page 2 - property and price.
heading(doc, "2. Property")
field(doc, "Address", "property_address")
field(doc, "Display address", "property_display_address")
field(doc, "Suburb", "property_suburb")
field(doc, "City", "property_city")
field(doc, "Property type", "property_type")
field(doc, "Erf number", "erf_number")
field(doc, "Unit / section", "property_unit_number")
field(doc, "Section number", "property_section_number")
field(doc, "Scheme / complex", "property_complex_name")
field(doc, "Estate", "property_estate_name")
field(doc, "Sectional title number", "sectional_title_number")
heading(doc, "3. Purchase price")
field(doc, "Purchase price", "purchase_price")
field(doc, "Deposit", "deposit_amount")
body(doc, "The Purchase Price is payable in accordance with the accepted offer, guarantees, bond approval, cash undertakings and conveyancer requirements.")

doc.add_page_break()
# Page 3 - finance and transfer.
heading(doc, "4. Finance")
field(doc, "Finance type", "finance_type")
field(doc, "Bond amount", "bond_amount")
field(doc, "Cash contribution", "cash_amount")
heading(doc, "5. Suspensive conditions")
body(doc, "{suspensive_conditions}")
body(doc, "If a suspensive condition is not fulfilled or waived within the agreed period, the parties must follow the consequence recorded in this agreement.")
heading(doc, "6. Occupation and transfer")
field(doc, "Occupation date", "occupation_date")
field(doc, "Expected transfer date", "transfer_date")
body(doc, "Risk, benefits and obligations transfer according to the final agreement terms and applicable conveyancing requirements.")

doc.add_page_break()
# Page 4 - fixtures, authority and commission.
heading(doc, "7. Fixtures and fittings")
body(doc, "The Property is sold together with fixtures and fittings of a permanent nature unless expressly excluded in Special Conditions or an annexure.")
heading(doc, "8. Capacity and authority")
body(doc, "Each party warrants that they have the necessary capacity and authority to sign this agreement. Legal entities must ensure the signatory is duly authorised.")
field(doc, "Seller representative", "seller_representative_name")
field(doc, "Representative capacity", "seller_representative_capacity")
field(doc, "Resolution date", "seller_resolution_date")
field(doc, "Authority basis", "seller_authority_basis")
heading(doc, "9. Commission")
field(doc, "Agency", "organisation_name")
field(doc, "Gross commission percentage", "gross_commission_percentage")
field(doc, "Gross commission amount", "gross_commission_amount")
field(doc, "Agency commission amount", "agency_commission_amount")
field(doc, "Agent commission amount", "agent_commission_amount")
body(doc, "Commission is earned and payable according to the accepted offer, mandate and applicable agency agreement.")

doc.add_page_break()
# Page 5 - special and general terms.
heading(doc, "10. Special conditions")
body(doc, "{special_conditions}")
field(doc, "Annexures", "annexures_list")
heading(doc, "11. General terms")
body(doc, "The parties choose their recorded addresses for notices and consent to the jurisdiction recorded in the final agreement. No amendment or cancellation is valid unless reduced to writing and signed or accepted by the parties as required.")
body(doc, "The parties consent to processing of personal information required for conveyancing, finance, verification, communication and transaction administration.")
body(doc, "The parties acknowledge that they have read, understood and accepted every page of this agreement and its annexures.")

doc.add_page_break()
# Page 6 - signing page aligned to the product's six-page OTP signing seed.
heading(doc, "12. Signatures")
body(doc, "Signed at ______________________________ on ______________________________.")
signature_block(doc, "Purchaser", "buyer_full_name")
signature_block(doc, "Seller", "seller_full_name")
heading(doc, "Witness", 2)
body(doc, "Name: ______________________________________________")
body(doc, "Signature: __________________________________________")
heading(doc, "Agency record", 2)
field(doc, "Agency", "organisation_name")
field(doc, "Agent", "agent_full_name")
field(doc, "FFC number", "agent_ffc_number")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
